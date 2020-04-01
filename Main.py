# !/usr/bin/python
# -*- coding: utf-8 -*-
import requests
import json
import os
from docx import Document
from datetime import date
import datetime
import calendar
import platform
import pprint
import configparser
from PIL import Image
import time
#Anterior API key:"eyJhbGciOiJIUzI1NiJ9.eyJzdWIiOiJhaHVydGFkb0B0dnJpb2phLmNvbSIsImp0aSI6IjBiYzQxMGQyLTc1NjAtNDYyMS05ZjgxLTEwOWE4N2YxZDE2YSIsImlzcyI6IkFFTUVUIiwiaWF0IjoxNTY1Njk1OTk2LCJ1c2VySWQiOiIwYmM0MTBkMi03NTYwLTQ2MjEtOWY4MS0xMDlhODdmMWQxNmEiLCJyb2xlIjoiIn0.kykNwBojwUwqO8r16SZ1NfdLxV2Bv-PR1PJUpBhxBMM"

print(
    'Script de generación automatica de predicción utilizando datos generados por la API de la AEMET, programado por Alejandro Hurtado, para TV Rioja')
webcams = ['logrono', 'haroo', 'calahorra', 'alfaro', 'cervera', 'najera', 'torrecilla', 'navarrete', 'stodomingo']
codigosmunicipio = {
    'calahorra': '26036',
    'logrono': '26089',
    'alfaro': '26011',
    'arnedo': '26018',
    'cervera': '26047',
    'ezcaray': '26061',
    'haro': '26071',
    'domingo': '26138',
    'torrecilla': '26151',
    'najera': '26102',
}


def descargarwebcams(nombrepoblacion):
    archivopoblacion = open('webcams/'+ nombrepoblacion + '.jpg', 'wb')
    urldewebcam = ('https://actualidad.larioja.org/images/webcam/' + nombrepoblacion + '.jpg')
    print('Descargando imagen de webcam desde ' + urldewebcam)
    archivopoblacion.write(requests.get(urldewebcam).content)
    archivopoblacion.close()
    convertirimagen = Image.open('webcams/' + nombrepoblacion + '.jpg')
    w, h =convertirimagen.size
    finalimagen=convertirimagen.crop((0,76, w, h))
    finalimagen = convertirimagen.resize((1920,1080))
    finalimagen.save('webcams/' + nombrepoblacion + '.png', 'PNG',optimize=True,quality=10)
    os.remove('webcams/' + nombrepoblacion + '.jpg')



def llamadaapipronostico(urldellamada, claveapi):
    payload = ""
    headers = {
        'accept': "application/json",
        'api_key': "eyJhbGciOiJIUzI1NiJ9.eyJzdWIiOiJ3ZWJAdHZyLmVzIiwianRpIjoiZmFkNGVlYmUtYzFlYi00ZjUwLWFkOGMtY2NlMTlkMDY4YjdhIiwiaXNzIjoiQUVNRVQiLCJpYXQiOjE1NjgzNjY2MTEsInVzZXJJZCI6ImZhZDRlZWJlLWMxZWItNGY1MC1hZDhjLWNjZTE5ZDA2OGI3YSIsInJvbGUiOiIifQ.ILskXBCqvM1uENPUkzHiFNF_0HcvQERcq_tnc-8p-uY"
    }
    respuesdeapi = requests.request("GET", urldellamada, data=payload, headers=headers)
    diccionarioderespuesta = json.loads(respuesdeapi.text)
    urlrespuesta = diccionarioderespuesta.get('datos', None)
    # Ahora este rollo porque los de la AEMET no saben utilizar correctamente JSON
    archivoderespuestalocal = open('temp.txt', 'wb')
    print("Descargando archivo de texto temporal desde " + urlrespuesta)
    archivoderespuestalocal.write(requests.get(urlrespuesta).content)
    archivoderespuestalocal.close()
    archivoderespuestalocal = open('temp.txt', 'rt', encoding='windows-1252')
    stringderespuesta = archivoderespuestalocal.read()
    archivoderespuestalocal.close()
    os.remove('temp.txt')
    stringderespuesta = (stringderespuesta.replace('\n', ' '))
    stringderespuesta = (stringderespuesta.translate({ord('\n'): None}))
    anterior, separador, prediccion = stringderespuesta.partition("B.- PREDICCIÓN")
    return prediccion


def creardocxpronostico(hoy, manana):
    documentodesalida = Document()
    nombredocumento = ('guiones/guion' + (str(date.today())) + '.docx')
    print('El nombre del docx que se creará como guión es ' + nombredocumento)
    documentodesalida.add_heading('GUIÓN VOZ EN OFF TIEMPO ' + str(date.today()), 0)
    documentodesalida.add_paragraph('Así ha amanecido en Logroño como se aprecia en este time-lapse de Meteo Sojuela.',
                                    style=None)
    documentodesalida.add_paragraph(hoy, style=None)
    documentodesalida.add_paragraph(manana, style=None)
    documentodesalida.add_paragraph(
        'Les dejamos con las imágenes que nos envían nuestros colaboradores del tiempo. \n\n\n')
    documentodesalida.save(nombredocumento)
    print(
        'Está usando un sistema ' + platform.system() + 'si no es Windows es probable que falle el intento de '
                                                        'impresión.')
    try:
        # os.startfile(nombredocumento, 'print')
        print('Aquí se imprimiría, pero de momento está desactivado para favorecer el debug.')
    except:
        print("Error de impresión")
    os.remove('Guion_pronosticos.txt')


def importarprediccionesespecificas(municipio, nombremunicipio):
    def interpretarcodigosestado(codigo):
        codigo=str(codigo)
        medionuboso = ['12', '13', '14', '15', '17']
        mediolluvia = ['23', '24', '25', '43', '44', '45']
        lluvia = ['26', '27', '46']
        nieve = ['33', '34', '35', '36']
        tormenta = ['51','52', '53', '54', '55', '56']
        if codigo == '11':
            #print('Despejado')
            return '1'
        elif codigo in medionuboso:
            #print('Medio nuboso')
            return '2'
        elif codigo == '16':
            #print('Cubierto')
            return '3'
        elif codigo in mediolluvia:
            #print('Medio lluvia')
            return '5'
        elif codigo in lluvia:
            #print('Lluvia')
            return '6'
        elif codigo in nieve:
            #print('Nieve')
            return '4'
        elif codigo in tormenta:
            #print('Tormenta')
            return '7'
        elif codigo == ' ':
            print("NO SE HAN RECIBIDO DATOS... INTENTELO MÁS TARDE PORQUE ES ERROR DE AEMET. SE DEVUELVE NUBLADO COMO PLACEHOLDER.")
            return '2'
        else:
            print(
                'NO SE RECONOCE EL CÓDIGO: ' + codigo + ' . SE VA A PONER EL CÓDIGO DE NUBLADO COMO PLACEHOLDER PERO HAY QUE REVISAR EL CODIGO RECIBIDO.')
            return '2'

    def interpretafecha(fecha):
        conversionfecha = datetime.datetime.strptime(fecha[0:10], '%Y-%m-%d').weekday()
        conversionfecha=(calendar.day_name[conversionfecha])
        if conversionfecha == 'Monday':
            return 1
        elif conversionfecha == 'Tuesday':
            return 2
        elif conversionfecha == 'Wednesday':
            return 3
        elif conversionfecha == 'Thursday':
            return 4
        elif conversionfecha == 'Friday':
            return 5
        elif conversionfecha == 'Saturday':
            return 6
        elif conversionfecha == 'Sunday':
            return 7

    payload = ''
    headers = {
        'accept': 'application/json',
        'api_key': 'eyJhbGciOiJIUzI1NiJ9.eyJzdWIiOiJ3ZWJAdHZyLmVzIiwianRpIjoiZmFkNGVlYmUtYzFlYi00ZjUwLWFkOGMtY2NlMTlkMDY4YjdhIiwiaXNzIjoiQUVNRVQiLCJpYXQiOjE1NjgzNjY2MTEsInVzZXJJZCI6ImZhZDRlZWJlLWMxZWItNGY1MC1hZDhjLWNjZTE5ZDA2OGI3YSIsInJvbGUiOiIifQ.ILskXBCqvM1uENPUkzHiFNF_0HcvQERcq_tnc-8p-uY'
    }
    urldeprediccion = ('https://opendata.aemet.es/opendata/api/prediccion/especifica/municipio/diaria/' + municipio)
    respuesdeapi = requests.request("GET", urldeprediccion, data=payload, headers=headers)
    diccionarioderespuesta = json.loads(respuesdeapi.text)
    urlrespuesta = diccionarioderespuesta.get('datos', None)
    archivojsonderespuesta = open('temp.json', 'wb')
    print("Descargando archivo de texto temporal desde " + urlrespuesta)
    archivojsonderespuesta.write(requests.get(urlrespuesta).content)
    archivojsonderespuesta.close()
    with open('temp.json', 'r', encoding='windows-1252') as archivojson:
        datos = archivojson.read()
    # A partir de aquí el código se vuelve raro, pero es fácil de entender. Adjunto documentación,
    # así que tranquilidad. Hago muchas variables y doy muchas vueltas porque he ido aprendiendo cómo iba su API a la
    # vez que programaba.  Quedará claro en la doc.
    diccionariodeprediccion = json.loads(datos)
    diossabequeestoyhaciendo = (diccionariodeprediccion[0])
    predicciondiccionario = diossabequeestoyhaciendo['prediccion']
    predicciondiccionario = predicciondiccionario['dia']
    diauno = predicciondiccionario[0]
    diados = predicciondiccionario[1]
    diatres = predicciondiccionario[2]
    diacuatro = predicciondiccionario[3]
    diacinco = predicciondiccionario[4]
    diaseis = predicciondiccionario[5]
    # Importa en variable la probabilidad de precipitación general del primer día
    precipitacionuno = diauno['probPrecipitacion']
    precipitacionuno = precipitacionuno[0]
    cantidadprecipitacionuno = precipitacionuno['value']

    # Importa en variable la probabilidad de precipitación general del segundo día
    precipitaciondos = diados['probPrecipitacion']
    precipitaciondos = precipitaciondos[0]
    cantidadprecipitaciondos = precipitaciondos['value']

    # Importa en variable la probabilidad de precipitación general del tercer día
    precipitaciontres = diatres['probPrecipitacion']
    precipitaciontres = precipitaciontres[0]
    cantidadprecipitaciontres = precipitaciontres['value']

    # Importa en variable la probabilidad de precipitación general del cuarto día
    precipitacioncuatro = diacuatro['probPrecipitacion']
    precipitacioncuatro = precipitacioncuatro[0]
    cantidadprecipitacioncuatro = precipitacioncuatro['value']
    # Importa en variable la probabilidad de precipitación general del quinto día
    precipitacioncinco = diacinco['probPrecipitacion']
    precipitacioncinco = precipitacioncinco[0]
    cantidadprecipitacioncinco = precipitacioncinco['value']
    #Sexto día
    precipitacionseis = diaseis['probPrecipitacion']
    precipitacionseis = precipitacionseis[0]
    cantidadprecipitacionseis = precipitacionseis['value']
    # Una vez importadas las probabilidades de lluvia, ahora iremos a las posibles temperaturas.
    tempuno = diauno['temperatura']
    tempmaxuno = tempuno['maxima']
    tempminuno = tempuno['minima']
    # La temperatura de mañana...
    tempdos = diados['temperatura']
    tempmaxdos = tempdos['maxima']
    tempmindos = tempdos['minima']

    # La temperatura de pasadomañana...
    temptres = diatres['temperatura']
    tempmaxtres = temptres['maxima']
    tempmintres = temptres['minima']

    # And so on...
    tempcuatro = diacuatro['temperatura']
    tempmaxcuatro = tempcuatro['maxima']
    tempmincuatro = tempcuatro['minima']
    # Por ultimo, la temperatura del dia cinco
    tempcinco = diacinco['temperatura']
    tempmaxcinco = tempcinco['maxima']
    tempmincinco = tempcinco['minima']
    #Temperaturas de sexto dia
    tempseis = diaseis['temperatura']
    tempmaxseis = tempseis['maxima']
    tempminseis = tempseis['minima']
    # Ahora el estado del cielo, que este tiene más intringulis. Primer estado
    estadouno = diauno['estadoCielo']
    estadouno = estadouno[0]
    estadouno = estadouno['value']
    estadouno = interpretarcodigosestado(estadouno)
    # Estado del cielo dos
    estadodos = diados['estadoCielo']
    estadodos = estadodos[0]
    estadodos = estadodos['value']
    estadodos = interpretarcodigosestado(estadodos)
    # Estado del cielo tres
    estadotres = diatres['estadoCielo']
    estadotres = estadotres[0]
    estadotres = estadotres['value']
    estadotres = interpretarcodigosestado(estadotres)
    # Estad del cuelo cuatro
    estadocuatro = diacuatro['estadoCielo']
    estadocuatro = estadocuatro[0]
    estadocuatro = estadocuatro['value']
    estadocuatro = interpretarcodigosestado(estadocuatro)
    # Estado del cielo cinco
    estadocinco = diacinco['estadoCielo']
    estadocinco = estadocinco[0]
    estadocinco = estadocinco['value']
    estadocinco = interpretarcodigosestado(estadocinco)
    #Estado del cielo seis
    estadoseis = diaseis['estadoCielo']
    estadoseis = estadoseis[0]
    estadoseis = estadoseis['value']
    estadoseis = interpretarcodigosestado(estadoseis)
    #Ahora cogeremos las fechas porque las necesitamos para escribir los dias de la semana
    fechauno = interpretafecha(diauno['fecha'])
    fechados = interpretafecha(diados['fecha'])
    fechatres = interpretafecha(diatres['fecha'])
    fechacuatro = interpretafecha(diacuatro['fecha'])
    fechacinco = interpretafecha(diacinco['fecha'])
    fechaseis = interpretafecha(diaseis['fecha'])
    # A ver ahora como escribo en un archivo ini toda esta basura
    configuracion = configparser.ConfigParser()
    configuracion['GENERAL'] = {'ciudad': nombremunicipio}
    configuracion['DIAUNO'] = {'tempmin': tempminuno,
                               'tempmax': tempmaxuno,
                               'estadocielo': estadouno,
                               'probprec': cantidadprecipitacionuno,
                               'fecha': fechauno}
    configuracion['DIADOS'] = {'tempmin': tempmindos,
                               'tempmax': tempmaxdos,
                               'estadocielo': estadodos,
                               'probprec': cantidadprecipitaciondos,
                               'fecha': fechados}
    configuracion['DIATRES'] = {'tempmin': tempmintres,
                                'tempmax': tempmaxtres,
                                'estadocielo': estadotres,
                                'probprec': cantidadprecipitaciontres,
                                'fecha' : fechatres}
    configuracion['DIACUATRO'] = {'tempmin': tempmincuatro,
                                  'tempmax': tempmaxcuatro,
                                  'estadocielo': estadocuatro,
                                  'probprec': cantidadprecipitacioncuatro,
                                  'fecha' : fechacuatro}
    configuracion['DIACINCO'] = {'tempmin': tempmincinco,
                                 'tempmax': tempmaxcinco,
                                 'estadocielo': estadocinco,
                                 'probprec': cantidadprecipitacioncinco,
                                 'fecha' : fechacinco}
    configuracion['DIASEIS'] = {'tempmin': tempminseis,
                                 'tempmax': tempmaxseis,
                                 'estadocielo': estadoseis,
                                 'probprec': cantidadprecipitacionseis,
                                 'fecha' : fechaseis}
    nombredearchivo = str(nombremunicipio + '.ini')
    with open('predicciones/' + nombredearchivo, 'w') as archivopredic:
        configuracion.write(archivopredic)
    os.remove('temp.json')

if not os.path.isdir('./webcams'):
        os.mkdir('webcams')
if not os.path.isdir('./guiones'):
        os.mkdir('guiones')
if not os.path.isdir('./predicciones'):
        os.mkdir('predicciones')
# Utilizando la función, descarga las imágenes de las webcam
for poblaciones in webcams:
    descargarwebcams(poblaciones)
# Ahora descargaremos en un archivo de texto las predicciones utilizando la API Open Data de AEMET
with open('Guion_pronosticos.txt', 'w') as archivopronosticos:
    archivopronosticos.write('Pronosticos\n')
    archivopronosticos.write('Así ha amanecido en Logroño como se aprecia en este time-lapse de Meteo Sojuela \n')
    pronosticohoy = llamadaapipronostico('https://opendata.aemet.es/opendata/api/prediccion/ccaa/hoy/rio', '')
    pronosticomanana = llamadaapipronostico('https://opendata.aemet.es/opendata/api/prediccion/ccaa/manana/rio', '')
    archivopronosticos.write('\nHoy,' + pronosticohoy)
    archivopronosticos.write('\nMañana,' + pronosticomanana)
    archivopronosticos.write('\nLes dejamos con las imágenes que nos envían nuestros colaboradores del tiempo')
creardocxpronostico(pronosticohoy, pronosticomanana)
importarprediccionesespecificas(codigosmunicipio['calahorra'], 'Calahorra')
importarprediccionesespecificas(codigosmunicipio['logrono'], 'Logroño')
importarprediccionesespecificas(codigosmunicipio['haro'], 'Haro')
importarprediccionesespecificas(codigosmunicipio['alfaro'], 'Alfaro')
importarprediccionesespecificas(codigosmunicipio['torrecilla'], 'Torrecilla')
importarprediccionesespecificas(codigosmunicipio['najera'], 'Najera')
importarprediccionesespecificas(codigosmunicipio['domingo'], 'StoDomingo')
importarprediccionesespecificas(codigosmunicipio['cervera'], 'Cervera')
importarprediccionesespecificas(codigosmunicipio['arnedo'], 'Arnedo')
importarprediccionesespecificas(codigosmunicipio['ezcaray'], 'Ezcaray')
print("SCRIPT TERMINADO. SE CERRARÁ SOLO EN TRES SEGUNDOS")
time.sleep(3)
